from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Response, Header, Depends
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from contextlib import asynccontextmanager
from dataclasses import dataclass
from typing import List, Optional
import pypdf
import docx
import io
import json
import asyncio
import logging
import requests
import numpy as np
import hashlib 
import uuid
import random
from datetime import datetime, timedelta, timezone
from sklearn.metrics.pairwise import haversine_distances, cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
from geopy.geocoders import Nominatim
from functools import lru_cache
from pydantic import BaseModel, Field
from google import genai
from google.genai import types

import os
from dotenv import load_dotenv
load_dotenv()

from backend.db import init_db, save_jd, get_all_jds, get_jd_by_id, update_jd, delete_jd, get_user_by_email, update_user_password


logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# --- CONFIGURATION ---
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")  
GEMINI_MODEL = "gemini-3.5-flash-lite" 
MAX_CONCURRENT = 2        

def get_genai_client():
    key = os.getenv("GEMINI_API_KEY", "")
    if not key:
        logger.error("GEMINI_API_KEY is missing from environment variables!")
        raise HTTPException(
            status_code=500, 
            detail="GEMINI_API_KEY environment variable is not configured. Please add GEMINI_API_KEY in your Vercel/Render Environment Variables settings."
        )
    return genai.Client(api_key=key, http_options=types.HttpOptions(timeout=60 * 1000))

# Global lazy client accessor
_client_instance = None
def get_client():
    global _client_instance
    if _client_instance is None:
        _client_instance = get_genai_client()
    return _client_instance

semaphore = asyncio.Semaphore(MAX_CONCURRENT)

# --- 1. THE HYBRID LOCATION ENGINE (FREE) ---
geolocator = Nominatim(user_agent="muuchstac_ats_hybrid")

@lru_cache(maxsize=1000)
def get_coordinates(city_name: str):
    if not city_name or city_name.lower() in ["not found", "n/a", "unknown"]:
        return None
    try:
        location = geolocator.geocode(f"{city_name}, India", timeout=5)
        if location:
            return (location.latitude, location.longitude, location.address)
    except Exception as e:
        logger.error(f"Geocoding failed for {city_name}: {e}")
    return None

def calculate_hybrid_location_score(candidate_city: str, target_city: str = "Borivali, Mumbai") -> dict:
    if not candidate_city or not target_city:
        return {"relevancy": "Unknown", "status": "Location unknown."}

    cand_clean = candidate_city.lower().replace("maharashtra", "").replace("india", "").strip(" ,.")
    target_clean = target_city.lower().replace("maharashtra", "").replace("india", "").strip(" ,.")
    
    cand_parts = [p.strip() for p in cand_clean.split(',') if p.strip()]
    target_parts = [p.strip() for p in target_clean.split(',') if p.strip()]

    if cand_clean == target_clean:
        return {"relevancy": "High", "status": "Exact match (Assumed Local)"}
    if len(cand_parts) == 1 and cand_parts[0] in target_parts:
        return {"relevancy": "High", "status": f"Broad city match - {cand_parts[0].title()} (Assumed Local)"}
    if len(target_parts) == 1 and target_parts[0] in cand_parts:
        return {"relevancy": "High", "status": f"Broad city match - {target_parts[0].title()} (Assumed Local)"}

    cand_data = get_coordinates(candidate_city)
    office_data = get_coordinates(target_city)

    if not cand_data or not office_data:
        return {"relevancy": "Unknown", "status": "Location unknown."}

    cand_coords = (cand_data[0], cand_data[1])
    office_coords = (office_data[0], office_data[1])
    cand_address = cand_data[2].lower() 

    cand_rad = np.radians([cand_coords])
    office_rad = np.radians([office_coords])
    
    dist_matrix = haversine_distances(office_rad, cand_rad)
    straight_line_km = dist_matrix[0][0] * 6371.0  
    
    if straight_line_km > 50.0:
        if "maharashtra" in cand_address or "maharashtra" in candidate_city.lower():
            return {"relevancy": "Low", "status": f"In-State ({straight_line_km:.1f} km away)"}
        else:
            return {"relevancy": "Relocation", "status": f"Out of State ({straight_line_km:.1f} km away)"}

    lon1, lat1 = office_coords[1], office_coords[0]
    lon2, lat2 = cand_coords[1], cand_coords[0]
    
    try:
        osrm_url = f"https://router.project-osrm.org/route/v1/driving/{lon1},{lat1};{lon2},{lat2}?overview=false"
        response = requests.get(osrm_url, timeout=3)
        if response.status_code == 200:
            data = response.json()
            if data.get("code") == "Ok":
                duration_seconds = data["routes"][0]["duration"]
                duration_mins = int(duration_seconds // 60)
                
                if duration_mins < 45:
                    return {"relevancy": "High", "status": f"~{duration_mins} min drive (Excellent Commute)"}
                elif duration_mins <= 90:
                    return {"relevancy": "Medium", "status": f"~{duration_mins} min drive (Moderate Commute)"}
                else:
                    return {"relevancy": "Low", "status": f"~{duration_mins} min drive (Tough Commute)"}
    except Exception as e:
        logger.warning(f"OSRM Routing failed (using haversine distance fallback): {e}")
    
    return {"relevancy": "Medium" if straight_line_km < 30 else "Low", "status": f"{straight_line_km:.1f} km (Straight Line)"}

# --- 1.5 HYBRID SEARCH & SEMANTIC VECTOR EMBEDDING ENGINE ---
SKILL_TAXONOMY = {
    "influencer scouting": ["creator outreach", "influencer sourcing", "creator discovery", "kol scouting", "influencer onboarding", "influencer database"],
    "creator outreach": ["influencer scouting", "creator sourcing", "kol outreach", "influencer communication", "brand outreach"],
    "content creation": ["video production", "reels creation", "shorts conceptualization", "scriptwriting", "content strategy", "storyboarding"],
    "script review": ["scriptwriting", "content editing", "storyboard adaptation", "copywriting"],
    "contract negotiation": ["commercial deal closure", "mou negotiation", "deliverable agreements", "compensation negotiation", "rate card negotiation"],
    "deal finalization": ["contract closure", "agreement signing", "influencer contracting", "vendor onboarding"],
    "engagement metrics": ["cpm", "cpa", "roas", "engagement rate", "click-through rate", "views", "reach"],
    "react": ["reactjs", "react.js", "frontend", "javascript framework"],
    "python": ["python3", "py", "fastapi", "django", "flask"],
    "meta ads": ["facebook ads", "fb ads", "instagram ads", "paid social"],
    "google ads": ["gads", "ppc", "sem", "google adwords"],
    "seo": ["search engine optimization", "organic traffic", "keyword research"],
    "sql": ["postgres", "postgresql", "mysql", "database management"]
}

def compute_tfidf_similarity(jd_text: str, resume_text: str) -> float:
    """Computes TF-IDF n-gram vector cosine similarity (0 to 100)."""
    if not jd_text.strip() or not resume_text.strip():
        return 0.0
    try:
        vectorizer = TfidfVectorizer(ngram_range=(1, 2), stop_words='english', sublinear_tf=True)
        matrix = vectorizer.fit_transform([jd_text, resume_text])
        sim = cosine_similarity(matrix[0:1], matrix[1:2])[0][0]
        return float(np.round(sim * 100.0, 1))
    except Exception as e:
        logger.error(f"TF-IDF similarity calculation error: {e}")
        return 0.0

def compute_gemini_embedding_similarity(jd_text: str, resume_text: str) -> float:
    """Computes dense semantic vector embedding similarity via Gemini gemini-embedding-001."""
    if not jd_text.strip() or not resume_text.strip():
        return 0.0
    try:
        res_jd = get_client().models.embed_content(
            model="gemini-embedding-001",
            contents=jd_text[:4000]
        )
        res_resume = get_client().models.embed_content(
            model="gemini-embedding-001",
            contents=resume_text[:4000]
        )
        
        vec_jd = np.array(res_jd.embeddings[0].values)
        vec_resume = np.array(res_resume.embeddings[0].values)
        
        norm_jd = np.linalg.norm(vec_jd)
        norm_resume = np.linalg.norm(vec_resume)
        
        if norm_jd == 0 or norm_resume == 0:
            return 0.0
            
        cosine_sim = np.dot(vec_jd, vec_resume) / (norm_jd * norm_resume)
        return float(np.round(max(0.0, cosine_sim) * 100.0, 1))
    except Exception as e:
        logger.warning(f"Gemini embedding API fallback to TF-IDF due to: {e}")
        return compute_tfidf_similarity(jd_text, resume_text)

def evaluate_hybrid_skills(required_skills_str: str, resume_text: str) -> dict:
    """Taxonomy-aware skill matching using synonym expansion graph."""
    if not required_skills_str.strip():
        return {"matched_skills": [], "missing_skills": [], "hybrid_match_pct": 100.0}

    req_skills = [s.strip().lower() for s in required_skills_str.split(",") if s.strip()]
    if not req_skills:
        return {"matched_skills": [], "missing_skills": [], "hybrid_match_pct": 100.0}

    resume_lower = resume_text.lower()
    matched = []
    missing = []

    for req in req_skills:
        if req in resume_lower:
            matched.append(req.title())
            continue
        
        synonyms = SKILL_TAXONOMY.get(req, [])
        found_synonym = False
        for syn in synonyms:
            if syn in resume_lower:
                matched.append(f"{req.title()} (via '{syn}')")
                found_synonym = True
                break
        
        if not found_synonym:
            missing.append(req.title())

    pct = (len(matched) / len(req_skills)) * 100.0
    return {
        "matched_skills": matched,
        "missing_skills": missing,
        "hybrid_match_pct": float(np.round(pct, 1))
    }

# --- 2. FASTAPI SETUP ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    init_db()
    yield

app = FastAPI(title="Muuchstac ATS Engine", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# --- IN-MEMORY SESSION & AUTH STORES ---
# session_token -> email
ACTIVE_SESSIONS = {}
# email -> {"code": str, "expiry": datetime}
RESET_CODES = {}

# Auth Schemas
class LoginRequest(BaseModel):
    email: str
    password: str

class ForgotPasswordRequest(BaseModel):
    email: str

class ResetPasswordRequest(BaseModel):
    email: str
    code: str
    new_password: str

# Auth Dependency
def check_auth(authorization: Optional[str] = Header(None)) -> str:
    """Verifies that the request has a valid Bearer token in the ACTIVE_SESSIONS store."""
    if not authorization:
        raise HTTPException(status_code=401, detail="Missing authorization header.")
    
    parts = authorization.split()
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise HTTPException(status_code=401, detail="Invalid authorization header format. Use 'Bearer <token>'.")
        
    token = parts[1]
    if token not in ACTIVE_SESSIONS:
        raise HTTPException(status_code=401, detail="Session expired or invalid. Please log in again.")
        
    return ACTIVE_SESSIONS[token]


@dataclass
class FilterConfig:
    min_exp: float
    max_exp: float  # 👉 NEW
    required_skills: str
    required_education: str
    target_loc: str
    mand_exp: bool
    mand_edu: bool
    mand_skill: bool
    mand_loc: bool
    passing_score: int
    role_preset: str = "general"

# 👉 NEW: Schema for JD Extraction
class JDExtractionAI(BaseModel):
    min_experience_years: float = Field(description="Minimum years of experience required (output 0 if not mentioned)")
    max_experience_years: float = Field(description="Maximum years of experience required (output 10 if not mentioned or if role is open-ended)") # 👉 NEW
    required_skills: str = Field(description="Top 5-7 core technical skills required, separated by commas")
    required_education: str = Field(description="The minimum education degree required (e.g., 'Bachelors', 'MBA'). Leave blank if not mentioned.")
    target_location: str = Field(description="The primary city or location for the job. Leave blank if remote/not mentioned.")

class CandidateEvaluationAI(BaseModel):
    candidate_name: str       = Field(default="Unknown Candidate", description="Full name extracted from resume")
    experience_score: int     = Field(default=0, description="Experience score out of 40")
    experience_details: str   = Field(default="", description="1 short sentence explaining why they got this experience score")
    skills_score: int         = Field(default=0, description="Skills score out of 30")
    skills_details: str       = Field(default="", description="1 short sentence explaining why they got this skills score")
    education_score: int      = Field(default=0, description="Education score out of 30")
    education_details: str    = Field(default="", description="1 short sentence explaining why they got this education score")
    score_justification: str  = Field(default="", description="One sentence overall justification")
    candidate_location: Optional[str] = Field(default="Unknown", description="Specific City or area where candidate lives (e.g., 'Navi Mumbai', 'Pune')")
    contact_email: Optional[str]      = Field(default="Not found", description="Candidate email")
    contact_phone: Optional[str]      = Field(default="Not found", description="Candidate phone")
    experience_years: float   = Field(default=0.0, description="Years of relevant work/practical experience")
    
    # Advanced Fresher & Role Specificity Fields
    full_time_years: float    = Field(default=0.0, description="Total full-time professional experience in years")
    internship_months: float  = Field(default=0.0, description="Total internship experience in months")
    equivalent_practical_years: float = Field(default=0.0, description="Weighted practical experience (Full-time yrs + (Internship months/12)*0.8 + Project credits)")
    candidate_type: str       = Field(default="Experienced Professional", description="Candidate type: 'Experienced Professional', 'High-Potential Intern / Fresher', or 'Entry-Level Candidate'")
    niche_fit_score: int      = Field(default=0, description="Alignment with D2C / Men's Grooming / FMCG / Beauty / Social Media niche out of 100")
    niche_fit_details: str    = Field(default="", description="Short explanation of candidate's D2C/Grooming niche alignment")
    role_fit_scouting: int    = Field(default=0, description="Fit score for Scouting role out of 100")
    role_fit_content: int     = Field(default=0, description="Fit score for Content Creation role out of 100")
    role_fit_finalization: int= Field(default=0, description="Fit score for Finalization role out of 100")

    # Key Work Experience Evidence
    top_deliverables: str     = Field(default="", description="1-sentence summary of candidate's single strongest work achievement or campaign result")
    work_evidence: List[str]  = Field(default_factory=list, description="At least 4 to 5 specific, bulleted key work experience evidence items extracted from resume (e.g. number of creators onboarded, campaign reach, scripts reviewed, contracts negotiated, internships completed, tools used)")

    skills: List[str]         = Field(default_factory=list, description="Top matching skills")
    missing_requirements: List[str] = Field(default_factory=list, description="Missing requirements")
    
    # Hybrid Search & Semantic Vector Engine Fields
    semantic_similarity_score: int = Field(default=0, description="Semantic Vector Similarity score (0 to 100) between candidate resume and JD")
    hybrid_skill_score: int       = Field(default=0, description="Taxonomy-aware skill match score (0 to 100)")

class CandidateEvaluation(CandidateEvaluationAI):
    total_score: int = 0
    location_relevancy: str = ""
    location_details: str = ""
    is_qualified: bool = True
    source_file: str = Field(default="")
    semantic_similarity_score: int = 0
    hybrid_skill_score: int = 0

import re

def ocr_document_with_gemini(file_bytes: bytes, filename: str) -> str:
    """Uses Gemini Multi-Modal Vision capability to extract text from scanned PDFs, images, or unreadable documents."""
    fn_lower = filename.lower()
    logger.info(f"File '{filename}' text stream is low/insufficient. Triggering Gemini Multi-Modal Vision OCR...")
    
    mime_type = "application/pdf"
    if fn_lower.endswith((".png", ".webp")):
        mime_type = "image/png"
    elif fn_lower.endswith((".jpg", ".jpeg")):
        mime_type = "image/jpeg"
        
    prompt = "Extract all text from this resume/CV document accurately. Preserve candidate name, contact details (email, phone, city), work experience, bullet points, skills, education, and dates."
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = get_client().models.generate_content(
                model=GEMINI_MODEL,
                contents=[
                    types.Part.from_bytes(data=file_bytes, mime_type=mime_type),
                    prompt
                ]
            )
            extracted = response.text.strip() if response and response.text else ""
            logger.info(f"Gemini Vision OCR successfully extracted {len(extracted)} characters from '{filename}'.")
            return extracted
        except Exception as e:
            err_str = str(e).lower()
            if ("429" in err_str or "resource_exhausted" in err_str or "quota" in err_str) and attempt < max_retries - 1:
                import time
                retry_match = re.search(r'retry in (\d+)', err_str) or re.search(r'retrydelay[\':\s]+(\d+)', err_str)
                wait_seconds = (int(retry_match.group(1)) + 2) if retry_match else ((attempt + 1) * 10)
                logger.warning(f"Gemini OCR API rate limit hit (429). Waiting {wait_seconds}s...")
                time.sleep(wait_seconds)
            else:
                logger.error(f"Gemini Vision OCR failed for '{filename}': {e}")
                return ""
    return ""

def is_text_sufficient(text: str) -> bool:
    """Checks if extracted text contains meaningful words rather than garbage encodings or watermarks."""
    if not text:
        return False
    words = re.findall(r'[a-zA-Z0-9]{2,}', text)
    alpha_chars = sum(c.isalpha() for c in text)
# --- 3. HELPER FUNCTIONS ---
def safe_float(val, default: float = 0.0) -> float:
    """Safely converts string or numeric values (e.g. '6.5 years', '6Y 6M', 6.5) to float without raising ValueError."""
    if val is None:
        return default
    if isinstance(val, (int, float)):
        return float(val)
    try:
        match = re.search(r'[-+]?\d*\.\d+|\d+', str(val))
        if match:
            return float(match.group())
        return default
    except Exception:
        return default

def safe_int(val, default: int = 0) -> int:
    """Safely converts string or numeric values to int without raising ValueError."""
    if val is None:
        return default
    if isinstance(val, int):
        return val
    try:
        return int(safe_float(val, float(default)))
    except Exception:
        return default

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    fn_lower = filename.lower()
    text = ""
    try:
        if fn_lower.endswith(".pdf"):
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                if getattr(reader, "is_encrypted", False):
                    try:
                        reader.decrypt("")
                    except Exception:
                        pass
                page_texts = []
                for p in reader.pages:
                    try:
                        t = p.extract_text()
                        if t: page_texts.append(t)
                    except Exception:
                        pass
                text = " ".join(page_texts)
            except Exception as pdf_err:
                logger.warning(f"pypdf failed for '{filename}': {pdf_err}")
                text = ""

            # 1. If text stream is missing, garbage, or insufficient, trigger Gemini Multi-Modal Vision OCR
            if not is_text_sufficient(text):
                ocr_text = ocr_document_with_gemini(file_bytes, filename)
                if is_text_sufficient(ocr_text):
                    text = ocr_text

            # 2. Raw ASCII stream extraction fallback if still insufficient
            if not is_text_sufficient(text):
                try:
                    raw_str = file_bytes.decode("latin-1", errors="ignore")
                    printable_blocks = re.findall(r'[a-zA-Z0-9\s.,@+\-:;()/]{4,}', raw_str)
                    candidate_text = " ".join(b.strip() for b in printable_blocks if len(b.strip()) > 3)
                    if is_text_sufficient(candidate_text):
                        text = candidate_text
                except Exception:
                    pass

            return text

        elif fn_lower.endswith(".docx"):
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text: parts.append(row_text)
                text = " ".join(parts)
            except Exception as docx_err:
                logger.warning(f"docx failed for '{filename}': {docx_err}")
                text = ""
                
            if not is_text_sufficient(text):
                text = ocr_document_with_gemini(file_bytes, filename)
            return text

        elif fn_lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
            # Image resume uploads (PNG / JPG / WEBP)
            return ocr_document_with_gemini(file_bytes, filename)

        elif fn_lower.endswith(".doc"):
            # Old binary MS Word 97-2003 format: extract printable strings + Gemini Vision fallback
            printable_strings = re.findall(r'[\x20-\x7E]{4,}', file_bytes.decode('latin-1', errors='ignore'))
            text = " ".join(printable_strings)
            if not is_text_sufficient(text):
                text = ocr_document_with_gemini(file_bytes, filename)
            return text

        else:
            return ""
    except Exception as e:
        logger.error(f"Error in extract_text_from_bytes for {filename}: {e}")
        return ""

def build_prompt(resume_text: str, jd: str, cfg: FilterConfig) -> str:
    skills_req = f"Required Skills: {cfg.required_skills} (Mandatory: {cfg.mand_skill})" if cfg.required_skills else ""
    edu_req = f"Required Education: {cfg.required_education} (Mandatory: {cfg.mand_edu})" if cfg.required_education else ""
    
    role_rubric = ""
    if cfg.role_preset == "scouting":
        role_rubric = """
SPECIAL RUBRIC - INFLUENCER SCOUTING (0-2 YRS):
- Focus Areas: Creator discovery across Instagram & YouTube, creator outreach, database management, engagement metrics evaluation, audience relevance, trend tracking.
- Freshers & Interns: Full credit for internships in influencer marketing, agency outreach, managing campus fest creators/ambassadors, running creator databases.
"""
    elif cfg.role_preset == "content":
        role_rubric = """
SPECIAL RUBRIC - INFLUENCER CONTENT CREATION (2 YRS FIXED):
- Focus Areas: Video conceptualization, script review/adaptation, video production coordination, campaign performance metrics (CPM, ROAS, Engagement Rate), brand narrative consistency.
- Candidate must show strong script breakdown and campaign analysis capability.
"""
    elif cfg.role_preset == "finalization":
        role_rubric = """
SPECIAL RUBRIC - INFLUENCER FINALIZATION (0-2 YRS):
- Focus Areas: Contract negotiation, commercial deal closing, MOUs, deliverable schedules, compensation negotiation, relationship management.
- Freshers & Interns: Give credit for commercial negotiation experience, agency account management internships, event/vendor deal closure.
"""

    return f"""You are an expert Talent Acquisition AI for Godrej Consumer Products Ltd (MUUCHSTAC - Men's Grooming). Respond in JSON matching CandidateEvaluationAI schema.

ROLE TYPE / PRESET: {cfg.role_preset.upper()}
JD: {jd}
RESUME: {resume_text}

REQUIREMENTS:
Experience Range: {cfg.min_exp} to {cfg.max_exp} yrs (Mandatory: {cfg.mand_exp})
{skills_req}
{edu_req}

{role_rubric}

CRITICAL RULES FOR FRESHER & INTERN EXPERIENCE EVALUATION:
1. Carefully extract `full_time_years` and `internship_months`.
2. Compute `equivalent_practical_years` = full_time_years + (internship_months / 12) * 0.8 + (0.25 if candidate managed personal creator channels/live projects else 0).
3. Set `experience_years` equal to `equivalent_practical_years` so that capable freshers with internship/project experience are NOT unfairly scored as 0 or rejected when min_exp is 0.
4. Classify `candidate_type` as 'Experienced Professional' (if full_time_years >= 1.5), 'High-Potential Intern / Fresher' (if internship_months >= 3 or project experience exists), or 'Entry-Level Candidate'.
5. Evaluate `niche_fit_score` (0-100) based on exposure to Men's Grooming, Skincare, D2C Brands, FMCG, Beauty, Instagram Reels, YouTube Shorts.
6. Provide score breakdowns for all 3 role fits (scouting, content, finalization).
7. DO NOT score Location. Just extract the precise city/neighborhood for candidate_location.
8. INFLUENCER MARKETING WORK EVIDENCE EXTRACTION:
Extract AT LEAST 4 TO 5 distinct, bulleted key work experience evidence points in `work_evidence` STRICTLY focusing on Influencer Marketing, Creator Collaborations, Campaign Execution, and Social Media Strategy.
Each bullet point MUST showcase specific influencer/social media proof of work from the resume, such as:
- Number of Instagram/YouTube creators scouted, contacted, or onboarded
- Influencer campaign performance metrics (e.g. views, engagement rate %, CPM, CPA, ROAS, conversions)
- Video content conceptualization, script review/editing for Reels/Shorts, or storyboards
- Rate card negotiations, commercial contract terms, deliverable MOUs, and compensation tracking
- Agency influencer internships, campus ambassador programs, or personal creator/social media channel management
Do NOT include generic or unrelated job duties. If candidate is a fresher, highlight their influencer agency internships, campus outreach projects, and social media/creator experience. Also summarize the single best influencer marketing highlight in `top_deliverables`.
"""

# 👉 NEW: The AI Router Function with Rate Limit Retry
async def call_ai_engine(prompt: str, ai_provider: str) -> dict:
    if ai_provider == "gemini":
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = await asyncio.to_thread(
                    get_client().models.generate_content,
                    model=GEMINI_MODEL,
                    contents=prompt,
                    config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=CandidateEvaluationAI)
                )
                return json.loads(response.text)
            except Exception as e:
                err_str = str(e).lower()
                if ("429" in err_str or "resource_exhausted" in err_str or "quota" in err_str) and attempt < max_retries - 1:
                    retry_match = re.search(r'retry in (\d+)', err_str) or re.search(r'retrydelay[\':\s]+(\d+)', err_str)
                    wait_seconds = (int(retry_match.group(1)) + 2) if retry_match else ((attempt + 1) * 10)
                    logger.warning(f"Gemini API rate limit hit (429). Waiting {wait_seconds}s for quota reset (Attempt {attempt + 1}/{max_retries})...")
                    await asyncio.sleep(wait_seconds)
                else:
                    raise e
        
    elif ai_provider == "openai":
        # Placeholder for future OpenAI integration
        raise NotImplementedError("OpenAI API is not yet configured.")
        
    elif ai_provider == "custom":
        # Placeholder for future Custom / Local API integration
        raise NotImplementedError("Custom API is not yet configured.")
        
    else:
        raise ValueError(f"Unknown AI Provider: {ai_provider}")

async def evaluate_resume(file_bytes: bytes, filename: str, jd: str, cfg: FilterConfig, ai_provider: str) -> dict:
    async with semaphore:
        clean_name = filename.rsplit('.', 1)[0].replace('_', ' ').replace('-', ' ').title()
        try:
            text = await asyncio.to_thread(extract_text_from_bytes, file_bytes, filename)
            
            if not text or len(text.strip()) < 20:
                logger.warning(f"File '{filename}': Could not extract readable text.")
                return {
                    "candidate_name": clean_name,
                    "experience_score": 0,
                    "experience_details": "Could not read text (scanned image or empty file)",
                    "skills_score": 0,
                    "skills_details": "No text extracted",
                    "education_score": 0,
                    "education_details": "No text extracted",
                    "score_justification": f"Rejected: Unreadable resume content in '{filename}' (Scanned image or unsupported encoding).",
                    "candidate_location": "Unknown",
                    "contact_email": "Not found",
                    "contact_phone": "Not found",
                    "experience_years": 0.0,
                    "full_time_years": 0.0,
                    "internship_months": 0.0,
                    "equivalent_practical_years": 0.0,
                    "candidate_type": "Entry-Level Candidate",
                    "niche_fit_score": 0,
                    "niche_fit_details": "Unreadable document",
                    "role_fit_scouting": 0,
                    "role_fit_content": 0,
                    "role_fit_finalization": 0,
                    "skills": [],
                    "missing_requirements": ["Readable text content"],
                    "total_score": 0,
                    "location_relevancy": "Unknown",
                    "location_details": "Location unknown",
                    "is_qualified": False,
                    "source_file": filename
                }

            prompt = build_prompt(text, jd, cfg)
            
            # 👉 Pass the prompt to the new router
            result = await call_ai_engine(prompt, ai_provider)
            result["contact_email"] = result.get("contact_email") or "Not found"
            result["contact_phone"] = result.get("contact_phone") or "Not found"
            result["candidate_name"] = result.get("candidate_name") or clean_name

            # --- HYBRID VECTOR EMBEDDING & TAXONOMY EVALUATION ---
            tfidf_sim = compute_tfidf_similarity(jd, text)
            gemini_sim = await asyncio.to_thread(compute_gemini_embedding_similarity, jd, text)
            hybrid_vector_sim = int(np.round(0.4 * tfidf_sim + 0.6 * gemini_sim))
            
            hybrid_skills_res = evaluate_hybrid_skills(cfg.required_skills, text)
            
            result["semantic_similarity_score"] = hybrid_vector_sim
            result["hybrid_skill_score"] = int(np.round(hybrid_skills_res["hybrid_match_pct"]))

            # Merge matched skills from taxonomy match
            if hybrid_skills_res["matched_skills"]:
                existing_skills = set(s.lower() for s in result.get("skills", []))
                for ms in hybrid_skills_res["matched_skills"]:
                    if ms.lower() not in existing_skills:
                        result.setdefault("skills", []).append(ms)

            # Automatic Python Experience Calculation Safeguard
            ft_yrs = safe_float(result.get("full_time_years"))
            intern_m = safe_float(result.get("internship_months"))
            epe_yrs = safe_float(result.get("equivalent_practical_years"))

            if epe_yrs <= 0:
                epe_yrs = round(ft_yrs + (intern_m / 12.0) * 0.8, 1)
                result["equivalent_practical_years"] = epe_yrs

            exp_years = safe_float(result.get("experience_years"))
            if exp_years <= 0 and epe_yrs > 0:
                exp_years = epe_yrs
                result["experience_years"] = exp_years

            is_qualified = True
            rejection_reasons = []
            
            cand_loc = str(result.get("candidate_location") or "Unknown")
            geo_data = await asyncio.to_thread(calculate_hybrid_location_score, cand_loc, cfg.target_loc)
            
            result["location_relevancy"] = geo_data["relevancy"]
            result["location_details"] = geo_data["status"]
            result["candidate_location"] = f'{cand_loc} - {geo_data["status"]}'

            # Mandatory Location filter
            if cfg.mand_loc and result["location_relevancy"] == "Relocation":
                is_qualified = False
                rejection_reasons.append(f"Location commute too far ({geo_data['status']})")

            # Mandatory Experience filter (uses experience_years / equivalent_practical_years)
            exp_years = safe_float(result.get("experience_years"))
            if cfg.mand_exp and (exp_years < cfg.min_exp or exp_years > cfg.max_exp):
                is_qualified = False
                rejection_reasons.append(f"Experience ({exp_years} yrs) outside required range ({cfg.min_exp}-{cfg.max_exp} yrs)")

            # Mandatory Skills filter
            skills_score = safe_int(result.get("skills_score"))
            if cfg.mand_skill and cfg.required_skills and skills_score == 0:
                is_qualified = False
                rejection_reasons.append(f"Missing mandatory skills ({cfg.required_skills})")

            # Mandatory Education filter
            edu_score = safe_int(result.get("education_score"))
            if cfg.mand_edu and cfg.required_education and edu_score == 0:
                is_qualified = False
                rejection_reasons.append(f"Missing mandatory education ({cfg.required_education})")

            exp_score = safe_int(result.get("experience_score"))
            calculated_total = exp_score + skills_score + edu_score
            result["total_score"] = calculated_total

            if calculated_total < cfg.passing_score:
                is_qualified = False
                rejection_reasons.append(f"Total score {calculated_total}/100 below threshold ({cfg.passing_score})")

            result["is_qualified"] = is_qualified
            if not is_qualified and rejection_reasons:
                result["score_justification"] = "Rejected: " + "; ".join(rejection_reasons) + "."

            result["source_file"] = filename          
            return result
            
        except Exception as exc:
            logger.error(f"Pipeline crashed for {filename}: {exc}")
            return {
                "candidate_name": clean_name,
                "experience_score": 0,
                "experience_details": f"Processing error: {str(exc)}",
                "skills_score": 0,
                "skills_details": "Processing error",
                "education_score": 0,
                "education_details": "Processing error",
                "score_justification": f"Rejected: Failed to evaluate '{filename}'. Error: {str(exc)}",
                "candidate_location": "Unknown",
                "contact_email": "Not found",
                "contact_phone": "Not found",
                "experience_years": 0.0,
                "full_time_years": 0.0,
                "internship_months": 0.0,
                "equivalent_practical_years": 0.0,
                "candidate_type": "Entry-Level Candidate",
                "niche_fit_score": 0,
                "niche_fit_details": "Processing error",
                "role_fit_scouting": 0,
                "role_fit_content": 0,
                "role_fit_finalization": 0,
                "skills": [],
                "missing_requirements": ["Valid resume document"],
                "total_score": 0,
                "location_relevancy": "Unknown",
                "location_details": "Location unknown",
                "is_qualified": False,
                "source_file": filename
            }

# --- 4. API ENDPOINTS ---

# --- AUTHENTICATION ENDPOINTS ---

@app.post("/api/auth/login")
async def auth_login(req: LoginRequest):
    user = await asyncio.to_thread(get_user_by_email, req.email)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid email or password.")
    
    hashed_input = hashlib.sha256(req.password.encode()).hexdigest()
    if user["password_hash"] != hashed_input:
        raise HTTPException(status_code=401, detail="Invalid email or password.")
        
    token = str(uuid.uuid4())
    ACTIVE_SESSIONS[token] = user["email"]
    return {"token": token, "email": user["email"]}


@app.post("/api/auth/logout")
async def auth_logout(authorization: Optional[str] = Header(None)):
    if not authorization:
        return {"status": "ok"}
    parts = authorization.split()
    if len(parts) == 2 and parts[0].lower() == "bearer":
        token = parts[1]
        ACTIVE_SESSIONS.pop(token, None)
    return {"status": "ok"}


@app.post("/api/auth/forgot-password")
async def auth_forgot_password(req: ForgotPasswordRequest):
    user = await asyncio.to_thread(get_user_by_email, req.email)
    if not user:
        raise HTTPException(status_code=404, detail="Email address not registered.")
        
    code = f"{random.randint(100000, 999999)}"
    expiry = datetime.now(timezone.utc) + timedelta(minutes=15)
    RESET_CODES[user["email"].lower().strip()] = {"code": code, "expiry": expiry}
    
    # Print code clearly to the server console log
    print("\n" + "="*50)
    print(" [AUTH SYSTEM] PASSWORD RESET CODE")
    print(f" User: {user['email']}")
    print(f" Verification Code: {code}")
    print(f" Expiration: 15 minutes")
    print("="*50 + "\n")
    logger.info(f"Password reset code generated for {user['email']} (printed in backend console)")
    
    return {"message": "Verification code logged in backend console."}


@app.post("/api/auth/reset-password")
async def auth_reset_password(req: ResetPasswordRequest):
    email_clean = req.email.lower().strip()
    if email_clean not in RESET_CODES:
        raise HTTPException(status_code=400, detail="No reset code requested for this email.")
        
    record = RESET_CODES[email_clean]
    if record["expiry"] < datetime.now(timezone.utc):
        RESET_CODES.pop(email_clean, None)
        raise HTTPException(status_code=400, detail="Verification code has expired. Please request a new one.")
        
    if record["code"] != req.code.strip():
        raise HTTPException(status_code=400, detail="Invalid verification code.")
        
    # Valid! Update the password
    new_hash = hashlib.sha256(req.new_password.encode()).hexdigest()
    success = await asyncio.to_thread(update_user_password, email_clean, new_hash)
    if not success:
        raise HTTPException(status_code=500, detail="Failed to update password in database.")
        
    # Clear code
    RESET_CODES.pop(email_clean, None)
    return {"message": "Password reset successful. Please log in with your new password."}


# --- JD DATABASE CRUD ENDPOINTS ---

class JDCreateRequest(BaseModel):
    title: str
    description: str
    skills: str = ""
    education: str = ""
    min_exp: float = 0
    max_exp: float = 7
    location: str = ""

class JDUpdateRequest(BaseModel):
    title: str
    description: str
    skills: str = ""
    education: str = ""
    min_exp: float = 0
    max_exp: float = 7
    location: str = ""

@app.get("/api/jds/")
async def list_jds(email: str = Depends(check_auth)):
    """List all saved Job Descriptions (summary: id, title, timestamps)."""
    try:
        return await asyncio.to_thread(get_all_jds)
    except Exception as e:
        logger.error(f"Failed to list JDs: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/jds/{jd_id}")
async def get_jd(jd_id: str, email: str = Depends(check_auth)):
    """Get full details of a saved Job Description by ID."""
    try:
        jd = await asyncio.to_thread(get_jd_by_id, jd_id)
        if not jd:
            raise HTTPException(status_code=404, detail="Job Description not found.")
        return jd
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get JD {jd_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/jds/")
async def create_jd(req: JDCreateRequest, email: str = Depends(check_auth)):
    """Save a new Job Description to the database."""
    if not req.title.strip():
        raise HTTPException(status_code=400, detail="Title is required.")
    if not req.description.strip():
        raise HTTPException(status_code=400, detail="Description is required.")
    try:
        return await asyncio.to_thread(
            save_jd, req.title.strip(), req.description.strip(),
            req.skills, req.education, req.min_exp, req.max_exp, req.location
        )
    except Exception as e:
        logger.error(f"Failed to create JD: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/jds/{jd_id}")
async def update_jd_endpoint(jd_id: str, req: JDUpdateRequest, email: str = Depends(check_auth)):
    """Update an existing Job Description."""
    if not req.title.strip():
        raise HTTPException(status_code=400, detail="Title is required.")
    try:
        result = await asyncio.to_thread(
            update_jd, jd_id, req.title.strip(), req.description.strip(),
            req.skills, req.education, req.min_exp, req.max_exp, req.location
        )
        if not result:
            raise HTTPException(status_code=404, detail="Job Description not found.")
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update JD {jd_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/jds/{jd_id}")
async def delete_jd_endpoint(jd_id: str, email: str = Depends(check_auth)):
    """Delete a Job Description from the database."""
    try:
        deleted = await asyncio.to_thread(delete_jd, jd_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="Job Description not found.")
        return {"status": "deleted", "id": jd_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete JD {jd_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- JD EXTRACTION & ANALYSIS ENDPOINTS ---

# 👉 The Endpoint that reads the JD
@app.post("/extract-jd-params/")
async def extract_jd_params(job_description: str = Form(...), ai_provider: str = Form("gemini"), email: str = Depends(check_auth)):
    prompt = f"""
    You are an expert technical recruiter. Read the following Job Description and extract the core requirements.
    Return the data exactly as requested in the JSON schema.
    
    JOB DESCRIPTION:
    {job_description}
    """
    
    try:
        # We reuse your existing AI Router to do the extraction!
        if ai_provider == "gemini":
            response = await asyncio.to_thread(
                get_client().models.generate_content,
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=JDExtractionAI)
            )
            return json.loads(response.text)
        else:
            raise NotImplementedError("Only Gemini is configured for JD extraction right now.")
    except Exception as e:
        logger.error(f"JD Extraction failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-batch-parallel/", response_model=list[CandidateEvaluation])
async def analyze_batch_parallel(
    response: Response,   
    files: List[UploadFile] = File(...),
    job_description: str = Form(...),
    min_experience_years: float = Form(0.0),
    max_experience_years: float = Form(7.0),
    mandatory_experience: bool = Form(False),
    required_skills: str = Form(""),
    mandatory_skills: bool = Form(False),
    required_education: str = Form(""),
    mandatory_education: bool = Form(False),
    target_location: str = Form("Borivali, Mumbai"),
    mandatory_location: bool = Form(False),
    passing_score: int = Form(75),
    shortlist_top_n: int = Form(0),
    role_preset: str = Form("general"), # 👉 NEW: Influencer role preset (scouting, content, finalization, general)
    ai_provider: str = Form("gemini"),
    email: str = Depends(check_auth)
):
    # Parse and prepare all valid resume files
    file_data = []
    for f in files:
        if f.filename.lower().endswith((".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".webp")):
            content = await f.read()
            file_data.append((content, f.filename))

    if not file_data: 
        raise HTTPException(status_code=400, detail="No valid resume files (.pdf, .docx, .doc, .png, .jpg, .jpeg) found.")

    cfg = FilterConfig(min_experience_years, max_experience_years, required_skills, required_education, target_location, 
                       mandatory_experience, mandatory_education, mandatory_skills, mandatory_location, passing_score, role_preset)

    # Evaluate all resumes in parallel
    tasks = [evaluate_resume(fb, fn, job_description, cfg, ai_provider) for fb, fn in file_data]
    raw_results = list(await asyncio.gather(*tasks))

    if not raw_results: 
        raise HTTPException(status_code=500, detail="All files failed to process.")

    # Sort results: qualified first, then by highest score descending

    raw_results.sort(key=lambda c: (not c.get("is_qualified", False), -c.get("total_score", 0)))
    return raw_results[:shortlist_top_n] if shortlist_top_n > 0 else raw_results

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")

@app.get("/")
async def read_index():
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))

app.mount("/", StaticFiles(directory=FRONTEND_DIR), name="static")